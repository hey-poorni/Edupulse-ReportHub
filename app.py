import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import logging
import io

# Set up logging
logging.basicConfig(level=logging.INFO)

# Streamlit Components
st.title("EDUPULSE REPORTHUB 📄")
st.subheader('Make Ur Work Simplified')

# Custom CSS for styling
st.markdown("""
    <style>
    h1 {
        color: green; 
        text-align: center; 
        margin-bottom: 0px;
    }
    h3 {
        text-align: center; 
        margin-top: 0;
    }
    .green-text {
        color: green;
    }
    .css-1j6homm label {
        color: green; 
        font-weight: bold;
    }
    .st-multi-select {
        background-color: #dfffdf;  /* Light green background */
    }
    .st-multi-select-option {
        color: green;  /* Green text for selected subjects */
    }
    </style>
    """, unsafe_allow_html=True)

# Sidebar for file upload and subject selection
with st.sidebar:
    linkedin_url = "https://www.linkedin.com/in/poorni22/"
    st.markdown(f"""
    <div class='green-text'>
    DEVELOPED BY:<br>
    <font color='white'>Shri Poornima A</font><br>
    </div>
    <a href='{linkedin_url}' style='color: green;'>LinkedIn</a>
    """, unsafe_allow_html=True)
    st.markdown("<hr>", unsafe_allow_html=True)

    # File upload section
    uploaded_file = st.file_uploader("CHOOSE A CSV FILE:", type=["csv"])

    # Generate Report Button
    generate_button = st.button("Generate Report") if uploaded_file else None

    # Subject selection from columns
    if uploaded_file is not None:
        dataframe = pd.read_csv(uploaded_file)
        non_subject_columns = [
            'sl.no', 'regno', 'name', 'student name', 'serial number',
            'registration number', 'Reg.No', 'student id', 'age', 'date', 'comments'
        ]
        subject_columns = [
            col for col in dataframe.columns
            if not any(term in col.lower() for term in non_subject_columns)
        ]
        selected_subjects = st.multiselect("SELECT SUBJECTS:", options=subject_columns)

# Function to calculate performance data
def calculate_performance(dataframe, selected_subjects, passing_mark=50):
    total_students = len(dataframe)
    subject_performance = {}

    subject_fail_counts = {col: 0 for col in selected_subjects}
    scores = {}
    failure_counts = [0] * (len(selected_subjects) + 1)

    name_column = next(
        (col for col in dataframe.columns if any(term in col.lower() for term in ['name', 'student'])),
        None
    )
    regno_column = next(
        (col for col in dataframe.columns if any(term in col.lower() for term in ['reg', 'registration', 'serial'])),
        None
    )

    for index, row in dataframe.iterrows():
        student_name = row.get(name_column, f'Student {index + 1}')
        reg_number = row.get(regno_column, 'Unknown')
        total_score = 0
        fail_subjects = 0

        for column in selected_subjects:
            subject_score = row[column] if pd.notna(row[column]) else 0
            subject_score_normalized = min(max(subject_score, 0), 100)
            total_score += subject_score_normalized

            if subject_score < passing_mark:
                subject_fail_counts[column] += 1
                fail_subjects += 1

        scores[student_name] = {'RegNo': str(reg_number), 'Marks': total_score}

        if 1 <= fail_subjects <= len(selected_subjects):
            failure_counts[fail_subjects - 1] += 1
        elif fail_subjects > len(selected_subjects):
            failure_counts[len(selected_subjects)] += 1

    sorted_scores = sorted(scores.items(), key=lambda x: x[1]['Marks'], reverse=True)
    top_5_scorers = sorted_scores[:5]
    least_5_scorers = sorted(scores.items(), key=lambda x: x[1]['Marks'])[:5]

    for column in selected_subjects:
        pass_count_subject = total_students - subject_fail_counts[column]
        subject_performance[column] = {
            'pass_count': pass_count_subject,
            'fail_count': subject_fail_counts[column],
            'pass_percentage': (pass_count_subject / total_students) * 100 if total_students > 0 else 0,
            'fail_percentage': (subject_fail_counts[column] / total_students) * 100 if total_students > 0 else 0,
        }

    overall_pass_percentage = (sum(stats['pass_count'] for stats in subject_performance.values()) / (total_students * len(subject_performance))) * 100 if total_students > 0 else 0
    overall_fail_percentage = (sum(stats['fail_count'] for stats in subject_performance.values()) / (total_students * len(subject_performance))) * 100 if total_students > 0 else 0

    return (subject_performance, overall_pass_percentage, overall_fail_percentage,
            top_5_scorers, least_5_scorers, failure_counts, name_column, regno_column)

# Function to set page borders and margins
def set_page_borders(doc):
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        sectPr = section._sectPr
        pgBorders = OxmlElement('w:pgBorders')
        for border_type in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_type}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '6')
            border.set(qn('w:space'), '24')
            border.set(qn('w:color'), 'GREEN')
            pgBorders.append(border)
        sectPr.append(pgBorders)

# Function to generate Word report
def generate_word_report(df, subject_performance, overall_pass_percentage, overall_fail_percentage,
                         top_5_scorers, least_5_scorers, failure_counts, name_column, regno_column):
    doc = Document()
    set_page_borders(doc)

    title = doc.add_paragraph()
    run = title.add_run("National Engineering college,k.R Nagar,Kovilparri-628503")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(18)
    run.bold = True
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    title = doc.add_paragraph()
    run = title.add_run("Student Performance Analysis Report")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.bold = True
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()

    overall_perf = doc.add_paragraph()
    overall_perf.add_run("OVERALL PERFORMANCE").bold = True
    overall_perf.add_run(f"\nPass Percentage: {overall_pass_percentage:.1f}%")
    overall_perf.add_run(f"\nFail Percentage: {overall_fail_percentage:.1f}%")
    doc.add_paragraph()

    multi_sub_fail = doc.add_paragraph()
    multi_sub_fail.add_run("MULTI-SUBJECT FAILURES COUNT:").bold = True
    for i in range(len(failure_counts) - 1):
        multi_sub_fail.add_run(f"\n{i + 1} Subject: {failure_counts[i]}")
    multi_sub_fail.add_run(f"\n{len(failure_counts)}+ Subjects: {failure_counts[-1]}")
    doc.add_paragraph()

    subject_perf = doc.add_paragraph()
    subject_perf.add_run("SUBJECT WISE PERFORMANCE:").bold = True

    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'SUBJECT'
    hdr_cells[1].text = 'PASS COUNT'
    hdr_cells[2].text = 'FAIL COUNT'
    hdr_cells[3].text = 'PASS (%)'
    hdr_cells[4].text = 'FAIL (%)'

    for subject, stats in subject_performance.items():
        if subject.lower() != 'reg.no':
            row_cells = table.add_row().cells
            row_cells[0].text = str(subject)
            row_cells[1].text = str(stats['pass_count'])
            row_cells[2].text = str(stats['fail_count'])
            row_cells[3].text = f"{stats['pass_percentage']:.1f}"
            row_cells[4].text = f"{stats['fail_percentage']:.1f}"

    # Top 5 Scorers Section
    doc.add_paragraph()
    top5 = doc.add_paragraph()
    top5.add_run("TOP 5 SCORERS:").bold = True
    top5_table = doc.add_table(rows=1, cols=3)
    hdr_cells = top5_table.rows[0].cells
    hdr_cells[0].text = 'RANK'
    hdr_cells[1].text = name_column
    hdr_cells[2].text = 'MARKS OBTAINED'

    for rank, (student, details) in enumerate(top_5_scorers, start=1):
        row_cells = top5_table.add_row().cells
        row_cells[0].text = str(rank)
        row_cells[1].text = str(student)
        row_cells[2].text = str(details['Marks'])

    # Least 5 Scorers Section
    doc.add_paragraph()
    least5 = doc.add_paragraph()
    least5.add_run("LEAST 5 SCORERS:").bold = True
    least5_table = doc.add_table(rows=1, cols=3)
    hdr_cells = least5_table.rows[0].cells
    hdr_cells[0].text = 'RANK'
    hdr_cells[1].text = name_column
    hdr_cells[2].text = 'MARKS OBTAINED'

    for rank, (student, details) in enumerate(least_5_scorers, start=1):
        row_cells = least5_table.add_row().cells
        row_cells[0].text = str(rank)
        row_cells[1].text = str(student)
        row_cells[2].text = str(details['Marks'])

    # Save the document to an in-memory file
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer

# Generate report on button click
if generate_button and uploaded_file and selected_subjects:
    subject_performance, overall_pass_percentage, overall_fail_percentage, top_5_scorers, least_5_scorers, failure_counts, name_column, regno_column = calculate_performance(dataframe, selected_subjects)
    buffer = generate_word_report(dataframe, subject_performance, overall_pass_percentage, overall_fail_percentage, top_5_scorers, least_5_scorers, failure_counts, name_column, regno_column)

    st.success('Report generated successfully!')
    st.download_button(
        label="Download Report",
        data=buffer,
        file_name="Student_Performance_Report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
